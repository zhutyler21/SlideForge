"""文档解析器 - 支持 Word/Markdown/纯文本（PDF 待开发）"""

from __future__ import annotations

from pathlib import Path


def _reject_doc(path: Path) -> str:
    raise ValueError("不支持旧版 .doc 格式，请先转换为 .docx")


def _reject_pdf(path: Path) -> str:
    raise ValueError("PDF 解析功能正在开发中，请先将 PDF 转换为 Word(.docx) 或 Markdown(.md) 格式")


# 待开发：图片格式支持
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff"}


def parse_document(file_path: str | Path) -> str:
    """
    自动识别文件类型并解析为纯文本

    Args:
        file_path: 文档路径

    Returns:
        解析后的文本内容

    Raises:
        ValueError: 不支持的文件格式
        FileNotFoundError: 文件不存在
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {path}")

    suffix = path.suffix.lower()

    if suffix in _IMAGE_EXTENSIONS:
        raise ValueError(
            f"图片格式 ({suffix}) 作为输入文档的功能正在开发中。\n"
            "目前图片仅支持作为「参考图片」用于风格迁移。"
        )

    parsers = {
        ".pdf": _reject_pdf,
        ".docx": _parse_docx,
        ".doc": _reject_doc,
        ".md": _parse_text,
        ".markdown": _parse_text,
        ".txt": _parse_text,
        ".text": _parse_text,
    }

    parser = parsers.get(suffix)
    if not parser:
        raise ValueError(
            f"不支持的文件格式: {suffix}\n"
            f"支持的格式: {', '.join(parsers.keys())}"
        )

    text = parser(path)
    if not text.strip():
        raise ValueError(f"文档内容为空: {path}")

    return text


def _parse_docx(path: Path) -> str:
    """使用 python-docx 解析 Word 文档"""
    from docx import Document

    doc = Document(str(path))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            # 保留标题层级信息
            if paragraph.style.name.startswith("Heading"):
                level = paragraph.style.name.replace("Heading ", "")
                prefix = "#" * int(level) if level.isdigit() else "#"
                parts.append(f"{prefix} {text}")
            else:
                parts.append(text)

    # 解析表格
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts)


def _parse_text(path: Path) -> str:
    """直接读取文本文件"""
    text = path.read_text(encoding="utf-8")

    # 去除 Markdown front matter (---...---)
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            text = text[end + 3:].strip()

    return text


def get_document_stats(text: str) -> dict[str, int]:
    """获取文档统计信息"""
    return {
        "char_count": len(text),
        "word_count": len(text.split()),
        "line_count": text.count("\n") + 1,
    }
